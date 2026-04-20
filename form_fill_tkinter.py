"""
form_fill_tkinter.py
--------------------
Generates a filled "Formulaire de Demande de Modification" Word document
from a product selected in a GUI, using data pulled from database.xlsx , I used database-sample.xlsx for confidentiality reasons but the structure is the same.

Requirements:
    pip install openpyxl python-docx

File structure expected (same folder as this script):
    database-sample.xlsx                            <- product database
    formulaire_de_demande_de_Modification.docx  <- Word template
    output/                                  <- generated documents saved here (auto-created)
"""

import os
from datetime import datetime

import openpyxl
import tkinter as tk
from tkinter import ttk, messagebox
from docx import Document
from typing import List, Dict, Tuple
from typing import Optional, Dict

# ---------------------------------------------------------------------------
# Configuration — edit these paths if your files are in a different location
# ---------------------------------------------------------------------------
BASE_DIR    = os.path.dirname(os.path.abspath(__file__))
EXCEL_PATH  = os.path.join(BASE_DIR, "database-sample.xlsx")
TEMPLATE    = os.path.join(BASE_DIR, "formulaire_de_demande_de_Modification.docx")
OUTPUT_DIR  = os.path.join(BASE_DIR, "output")


# ---------------------------------------------------------------------------
# Data layer
# ---------------------------------------------------------------------------

def load_products(excel_path: str) -> Tuple[List, Dict]:
    """
    Read database-sample.xlsx and return:
        products       : ordered list of unique product names (column A)
        product_dosages: {product_name: [dosage, ...]} (column D)
    """
    wb = openpyxl.load_workbook(excel_path)
    sheet = wb.active

    products = []
    product_dosages = {}

    for row in range(2, sheet.max_row + 1):
        name   = sheet[f"A{row}"].value
        dosage = sheet[f"D{row}"].value

        name   = str(name).strip()   if name   else ""
        dosage = str(dosage).strip() if dosage else ""

        if not name:
            continue

        if name not in products:
            products.append(name)
            product_dosages[name] = []

        if dosage and dosage not in product_dosages[name]:
            product_dosages[name].append(dosage)

    return products, product_dosages

def get_product_data(excel_path: str, product_name: str, dosage: str) -> Optional[Dict]:
    """
    Return a dict of placeholder → value for the selected product/dosage row.
    Column mapping mirrors the placeholders used in the Word template.
    """
    wb = openpyxl.load_workbook(excel_path)
    sheet = wb.active

    # Placeholder codes used inside the Word template
    col_map = {
        "A": "@1",     # Nom commercial
        "B": "@2",     # DCI
        "C": "@3",     # Forme pharmaceutique
        "D": "@4",     # Dosage
        "E": "@5",     # Type de conditionnement
        "F": "@6",     # Prémix ou non
        "G": "@+12",   # Code ATC
        "H": "@+13",   # Classe pharmaco-thérapeutique
        "I": "@+14",   # Indication(s) thérapeutique(s)
        "K": "@+16",   # Raison sociale fabricant PA
        "L": "@+17",   # Adresse fabricant PA
        "M": "@+17,5", # Adresse fabricant prémix
        "N": "@+20",   # Numéro d'ordre DE
        "O": "@+21",   # Numéro d'identification administrative
        "P": "@+22",   # CNDT
    }

    for row in range(2, sheet.max_row + 1):
        row_name   = sheet[f"A{row}"].value
        row_dosage = sheet[f"D{row}"].value

        row_name   = str(row_name).strip()   if row_name   else ""
        row_dosage = str(row_dosage).strip() if row_dosage else ""

        if row_name.lower() == product_name.lower() and row_dosage.lower() == dosage.lower():
            data = {}
            for col, placeholder in col_map.items():
                value = sheet[f"{col}{row}"].value
                if value is None:
                    value = ""
                if placeholder == "@+14":
                    value = str(value).replace("<br>", "\n")
                data[placeholder] = str(value)
            return data

    return None


# ---------------------------------------------------------------------------
# Document generation
# ---------------------------------------------------------------------------

def fill_document(template_path: str, output_path: str, replacements: dict) -> None:
    """
    Load the Word template, replace all placeholders, and save to output_path.
    Handles both paragraphs and table cells.
    """
    doc = Document(template_path)

    def replace_in_element(element):
        """Recursively replace placeholders in paragraphs and tables."""
        if hasattr(element, "paragraphs"):
            for paragraph in element.paragraphs:
                full_text = "".join(run.text for run in paragraph.runs)
                new_text  = full_text
                for key, value in replacements.items():
                    new_text = new_text.replace(key, value)
                if new_text != full_text:
                    for run in paragraph.runs:
                        run.text = ""
                    if paragraph.runs:
                        paragraph.runs[0].text = new_text
                    else:
                        paragraph.add_run(new_text)

        if hasattr(element, "tables"):
            for table in element.tables:
                for row in table.rows:
                    for cell in row.cells:
                        replace_in_element(cell)

    replace_in_element(doc)
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)


# ---------------------------------------------------------------------------
# GUI
# ---------------------------------------------------------------------------

class ModificationFormApp:

    def __init__(self, root: tk.Tk):
        self.root = root
        self.root.title("Formulaire de Demande de Modification")
        self.root.geometry("800x620")
        self.root.resizable(False, False)

        # Check required files exist before building the UI
        for path, label in [(EXCEL_PATH, "database-sample.xlsx"), (TEMPLATE, "Word template")]:
            if not os.path.exists(path):
                messagebox.showerror(
                    "Fichier manquant",
                    f"{label} introuvable :\n{path}\n\n"
                    "Placez le fichier dans le même dossier que ce script."
                )
                root.destroy()
                return

        self.products, self.product_dosages = load_products(EXCEL_PATH)

        self._build_ui()

    # ------------------------------------------------------------------
    def _build_ui(self):
        frame = ttk.Frame(self.root, padding=24)
        frame.grid(row=0, column=0, sticky="nsew")
        self.root.columnconfigure(0, weight=1)
        self.root.rowconfigure(0, weight=1)
        frame.columnconfigure(1, weight=1)

        ttk.Label(
            frame,
            text="Formulaire de Demande de Modification",
            font=("Arial", 15, "bold")
        ).grid(row=0, column=0, columnspan=2, pady=(0, 20))

        # --- Fields pulled from the database ---
        ttk.Label(frame, text="Produit commercial :").grid(row=1, column=0, sticky="w", pady=6)
        self.product_var = tk.StringVar()
        self.product_combo = ttk.Combobox(
            frame, textvariable=self.product_var,
            values=self.products, state="readonly", width=50
        )
        self.product_combo.grid(row=1, column=1, sticky="ew", pady=6, padx=(12, 0))
        self.product_combo.bind("<<ComboboxSelected>>", self._on_product_selected)

        ttk.Label(frame, text="Dosage / Conditionnement :").grid(row=2, column=0, sticky="w", pady=6)
        self.dosage_var = tk.StringVar()
        self.dosage_combo = ttk.Combobox(
            frame, textvariable=self.dosage_var, state="readonly", width=50
        )
        self.dosage_combo.grid(row=2, column=1, sticky="ew", pady=6, padx=(12, 0))

        ttk.Separator(frame, orient="horizontal").grid(
            row=3, column=0, columnspan=2, sticky="ew", pady=12
        )

        # --- Fields entered manually ---
        manual_fields = [
            ("Nature de la modification :",      "nature_var",          ["administrative", "qualitative", "sécurité"], "combo"),
            ("Type de modification :",           "type_mod_var",        ["majeure", "mineure"],                        "combo"),
            ("Identification de la modification :", "identification_var", None, "entry"),
            ("N° et date du justificatif de paiement :", "justificatif_var", None, "entry"),
            ("N° et date du bordereau de versement :",   "bordereau_var",    None, "entry"),
            ("Taux d'intégration :",             "taux_var",            None, "entry"),
            ("N° et date du ML :",               "ml_var",              None, "entry"),
            ("N° et date du GMP :",              "gmp_var",             None, "entry"),
        ]

        self.field_widgets = {}
        for i, (label, attr, values, kind) in enumerate(manual_fields, start=4):
            ttk.Label(frame, text=label).grid(row=i, column=0, sticky="w", pady=4)
            var = tk.StringVar()
            setattr(self, attr, var)
            if kind == "combo":
                widget = ttk.Combobox(frame, textvariable=var, values=values, state="readonly", width=50)
            else:
                widget = ttk.Entry(frame, textvariable=var, width=52)
            widget.grid(row=i, column=1, sticky="ew", pady=4, padx=(12, 0))
            self.field_widgets[attr] = widget

        ttk.Button(
            frame,
            text="  Générer le document  ",
            command=self._generate
        ).grid(row=len(manual_fields) + 4, column=0, columnspan=2, pady=22)

    # ------------------------------------------------------------------
    def _on_product_selected(self, _event=None):
        product = self.product_var.get()
        dosages = self.product_dosages.get(product, [])
        self.dosage_combo["values"] = dosages
        if len(dosages) == 1:
            self.dosage_var.set(dosages[0])
        else:
            self.dosage_var.set("")

    # ------------------------------------------------------------------
    def _generate(self):
        # Validate required fields
        required = {
            "Produit":              self.product_var.get(),
            "Dosage":               self.dosage_var.get(),
            "Nature":               self.nature_var.get(),
            "Type de modification": self.type_mod_var.get(),
        }
        for label, value in required.items():
            if not value:
                messagebox.showerror("Champ manquant", f"Veuillez renseigner : {label}")
                return

        # Fetch database row
        excel_data = get_product_data(EXCEL_PATH, self.product_var.get(), self.dosage_var.get())
        if excel_data is None:
            messagebox.showerror(
                "Erreur",
                f"Produit introuvable dans la base de données :\n"
                f"{self.product_var.get()} — {self.dosage_var.get()}"
            )
            return

        # Manual fields
        user_data = {
            "@7":   self.nature_var.get(),
            "@8":   self.identification_var.get(),
            "@9":   self.type_mod_var.get(),
            "@+10": self.justificatif_var.get(),
            "@+11": self.bordereau_var.get(),
            "@+15": self.taux_var.get(),
            "@+18": self.ml_var.get(),
            "@+19": self.gmp_var.get(),
            "@+24": datetime.now().strftime("%d/%m/%Y"),
        }

        replacements = {**excel_data, **user_data}

        # Output path
        safe_name   = self.product_var.get().replace(" ", "_")
        output_path = os.path.join(OUTPUT_DIR, f"Formulaire_{safe_name}.docx")

        try:
            fill_document(TEMPLATE, output_path, replacements)
            messagebox.showinfo(
                "Succès",
                f"Document généré avec succès :\n{output_path}"
            )
        except Exception as exc:
            messagebox.showerror("Erreur", f"Une erreur est survenue :\n{exc}")


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

def main():
    root = tk.Tk()
    ModificationFormApp(root)
    root.mainloop()


if __name__ == "__main__":
    main()
